import io
import base64
from pathlib import Path

from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from pydantic import BaseModel

import pdfplumber
from docx import Document
from docx.shared import RGBColor
import fitz  # PyMuPDF

MAX_UPLOAD_BYTES = 50 * 1024 * 1024  # 50 MB

app = FastAPI(title="PII Redactor")

frontend_path = Path(__file__).parent.parent / "frontend"

# ---------------------------------------------------------------------------
# Lazy-load the model so startup is fast
# ---------------------------------------------------------------------------
_classifier = None

def get_classifier():
    global _classifier
    if _classifier is None:
        from transformers import pipeline
        print("Loading openai/privacy-filter model…")
        _classifier = pipeline(
            task="token-classification",
            model="openai/privacy-filter",
            aggregation_strategy="simple",
        )
        print("Model loaded.")
    return _classifier


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

LABEL_COLORS = {
    "private_person":  "#FF6B6B",
    "private_email":   "#4ECDC4",
    "private_phone":   "#FFE66D",
    "private_address": "#A8E6CF",
    "private_url":     "#C3B1E1",
    "private_date":    "#FFB347",
    "account_number":  "#87CEEB",
    "secret":          "#FF8C94",
}

def label_color(label: str) -> str:
    clean = label.split("-")[-1]
    return LABEL_COLORS.get(clean, "#DDDDDD")


def run_model(text: str) -> list[dict]:
    clf = get_classifier()
    results = clf(text)
    spans = []
    for r in results:
        spans.append({
            "start": r["start"],
            "end": r["end"],
            "word": r["word"],
            "label": r["entity_group"],
            "score": round(float(r["score"]), 3),
            "color": label_color(r["entity_group"]),
            "accepted": True,
        })
    return spans


def extract_text_from_pdf(data: bytes) -> tuple[str, list[dict]]:
    """Extract text and per-word bounding boxes from a PDF.

    Returns (full_text, word_coords). word_coords maps character positions
    in full_text back to page coordinates for precise, position-aware redaction.
    """
    parts = []
    word_coords = []
    offset = 0

    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page_num, page in enumerate(pdf.pages):
            words = page.extract_words()
            for i, w in enumerate(words):
                word_text = w["text"]
                word_len = len(word_text)
                word_coords.append({
                    "char_start": offset,
                    "char_end": offset + word_len,
                    "page": page_num,
                    "x0": float(w["x0"]),
                    "top": float(w["top"]),
                    "x1": float(w["x1"]),
                    "bottom": float(w["bottom"]),
                })
                parts.append(word_text)
                offset += word_len
                if i < len(words) - 1:
                    parts.append(" ")
                    offset += 1
            if words:
                parts.append("\n")
                offset += 1

    return "".join(parts), word_coords


def extract_text_from_docx(data: bytes) -> tuple[str, list[str]]:
    """Extract text from body paragraphs, tables, and headers/footers."""
    doc = Document(io.BytesIO(data))
    parts = []

    for section in doc.sections:
        for para in section.header.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for para in section.footer.paragraphs:
            if para.text.strip():
                parts.append(para.text)

    for para in doc.paragraphs:
        parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        parts.append(para.text)

    return "\n".join(parts), parts


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.get("/")
def index():
    return FileResponse(str(frontend_path / "index.html"))


class RedactRequest(BaseModel):
    text: str
    spans: list[dict]
    file_type: str
    original_b64: str
    word_coords: list[dict] = []


@app.post("/api/analyze")
async def analyze(file: UploadFile = File(...)):
    data = await file.read()
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(413, "Filen är för stor (max 50 MB)")

    filename = file.filename or ""
    ext = filename.rsplit(".", 1)[-1].lower()

    try:
        if ext == "pdf":
            text, word_coords = extract_text_from_pdf(data)
            file_type = "pdf"
        elif ext in ("docx", "doc"):
            text, _ = extract_text_from_docx(data)
            word_coords = []
            file_type = "docx"
        else:
            raise HTTPException(400, "Stöder bara PDF och Word (.docx)")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(422, f"Kunde inte läsa filen: {e}")

    try:
        spans = run_model(text)
    except Exception as e:
        raise HTTPException(500, f"Modellfel: {e}")

    b64 = base64.b64encode(data).decode()

    return JSONResponse({
        "text": text,
        "spans": spans,
        "file_type": file_type,
        "original_b64": b64,
        "filename": filename,
        "word_coords": word_coords,
    })


@app.post("/api/redact")
async def redact(req: RedactRequest):
    try:
        raw = base64.b64decode(req.original_b64)
    except Exception:
        raise HTTPException(400, "Ogiltig fildata")

    accepted_spans = [s for s in req.spans if s.get("accepted", True)]

    try:
        if req.file_type == "pdf":
            redacted = redact_pdf(raw, accepted_spans, req.word_coords)
            return StreamingResponse(
                io.BytesIO(redacted),
                media_type="application/pdf",
                headers={"Content-Disposition": "attachment; filename=redacted.pdf"},
            )
        else:
            redacted = redact_docx(raw, accepted_spans)
            return StreamingResponse(
                io.BytesIO(redacted),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": "attachment; filename=redacted.docx"},
            )
    except Exception as e:
        raise HTTPException(500, f"Maskeringsfel: {e}")


def redact_pdf(data: bytes, spans: list[dict], word_coords: list[dict]) -> bytes:
    """Coordinate-based PDF redaction: only redacts the exact positions the model flagged."""
    doc = fitz.open(stream=data, filetype="pdf")

    for span in spans:
        s, e = span["start"], span["end"]
        matches = [w for w in word_coords if w["char_start"] < e and w["char_end"] > s]
        for w in matches:
            page = doc[w["page"]]
            # +1 pt padding ensures glyph edges are covered
            rect = fitz.Rect(w["x0"] - 1, w["top"] - 1, w["x1"] + 1, w["bottom"] + 1)
            page.add_redact_annot(rect, fill=(0, 0, 0))

    for page in doc:
        page.apply_redactions()

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _redact_runs(runs, words: set) -> None:
    for run in runs:
        for word in words:
            if word in run.text:
                run.text = run.text.replace(word, "█" * len(word))
                run.font.color.rgb = RGBColor(0, 0, 0)


def redact_docx(data: bytes, spans: list[dict]) -> bytes:
    doc = Document(io.BytesIO(data))
    words = {s["word"] for s in spans}

    for para in doc.paragraphs:
        _redact_runs(para.runs, words)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _redact_runs(para.runs, words)

    for section in doc.sections:
        for para in section.header.paragraphs:
            _redact_runs(para.runs, words)
        for para in section.footer.paragraphs:
            _redact_runs(para.runs, words)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
